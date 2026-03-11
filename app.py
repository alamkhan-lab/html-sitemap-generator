import streamlit as st
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from urllib.parse import urlparse
import io
import re

# --- UI CONFIG ---
st.set_page_config(page_title="Sitemap Architect Pro", page_icon="🏗️")

# --- HYPERLINK HELPER ---
def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# --- SMART LABEL EXTRACTION ---
def get_meaningful_label(url):
    """
    Analyzes URL segments to find the most valuable anchor text.
    Ignores IDs like '12063' and category codes like 'c'.
    """
    path = urlparse(url).path.strip('/')
    if not path: return "Home"
    
    segments = path.split('/')
    # List of segments to ignore (common e-comm markers)
    ignore_list = ['c', 'p', 'v', 'cat', 'products', 'collections', 'category']
    
    # Iterate backwards through segments to find the first descriptive word
    for seg in reversed(segments):
        # 1. Skip if the segment is purely numeric (e.g., '12063')
        if seg.isdigit():
            continue
        # 2. Skip if the segment is in our ignore list (e.g., 'c')
        if seg.lower() in ignore_list:
            continue
        # 3. Clean and return the first valid word found
        label = seg.replace('-', ' ').replace('_', ' ').strip()
        return label.title()
        
    return "Link"

def get_smart_cluster(label):
    label_up = label.upper()
    clusters = ['LOAN', 'SUNSCREEN', 'SERUM', 'COSMETICS', 'SHAMPOO', 'SKINCARE', 'LIPSTICK', 'MAKEUP', 'BAGS', 'FASHION']
    for c in clusters:
        if c in label_up: return c.title()
    words = label.split()
    return f"{words[0].title()} {words[1].title()}" if len(words) >= 2 else label.title()

# --- AGGRESSIVE URL EXTRACTION ---
def extract_urls_robust(xml_content):
    """Uses both BeautifulSoup and Regex to find <loc> tags in XML."""
    if isinstance(xml_content, bytes):
        xml_content = xml_content.decode('utf-8', errors='ignore')
    
    # Method 1: BS4
    soup = BeautifulSoup(xml_content, 'xml')
    urls = [loc.text.strip() for loc in soup.find_all('loc')]
    
    # Method 2: Regex fallback (if BS4 fails due to namespaces)
    if not urls:
        urls = re.findall(r'<loc>(.*?)</loc>', xml_content)
    
    return list(set(urls))

def get_urls_via_request(url):
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36'}
    try:
        response = requests.get(url, headers=headers, timeout=15)
        if response.status_code == 200:
            return extract_urls_robust(response.content)
        elif response.status_code == 403:
            st.error("🚫 Nykaa/Big sites block automated tools. Please use the 'Paste Raw XML' tab.")
    except Exception as e:
        st.error(f"Error: {e}")
    return []

# --- ORGANIZATION ---
def organize_urls(urls):
    tree = {}
    for url in urls:
        parsed = urlparse(url)
        path_segments = [s for s in parsed.path.split('/') if s]
        
        # High Level Folder
        section = path_segments[0].title() if path_segments else "Main"
        
        # Meaningful Anchor Text
        page_label = get_meaningful_label(url)
        
        # Group Subject
        group = get_smart_cluster(page_label)

        if section not in tree: tree[section] = {}
        if group not in tree[section]: tree[section][group] = []
        tree[section][group].append((page_label, url))
    return tree

def create_docx(tree, domain):
    doc = Document()
    doc.add_heading(f'Sitemap Report: {domain}', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    for section in sorted(tree.keys()):
        doc.add_heading(section, level=1)
        for group, links in tree[section].items():
            if len(links) > 1: doc.add_heading(group, level=2)
            for text, link in links:
                p = doc.add_paragraph(style='List Bullet')
                add_hyperlink(p, text, link)
    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    return bio

# --- UI ---
st.title("🏗️ Sitemap Architect Pro")
st.write("Specialized engine for e-commerce (Nykaa-style) and BFSI sitemaps.")

tab1, tab2, tab3 = st.tabs(["🔗 XML URL", "📄 Paste Raw XML", "📝 Naked URLs"])

final_urls = []
domain_name = "Website"

with tab1:
    xml_url = st.text_input("Enter Sitemap URL")

with tab2:
    raw_xml_content = st.text_area("Paste the XML Code here (Ctrl+A / Ctrl+V from browser)", height=200)

with tab3:
    naked_urls = st.text_area("Paste Naked URL list", height=200)

if st.button("Generate Structured Sitemap"):
    if xml_url:
        domain_name = urlparse(xml_url).netloc
        final_urls = get_urls_via_request(xml_url)
    
    if not final_urls and raw_xml_content:
        domain_name = "Manual-Entry"
        final_urls = extract_urls_robust(raw_xml_content)
        
    if not final_urls and naked_urls:
        final_urls = [l.strip() for l in naked_urls.split('\n') if l.strip()]
        domain_name = urlparse(final_urls[0]).netloc if final_urls else "List"

    if final_urls:
        if len(final_urls) > 3000:
            st.warning("Large sitemap. Capping at 3000 links.")
            final_urls = final_urls[:3000]
            
        tree = organize_urls(final_urls)
        docx_data = create_docx(tree, domain_name)
        
        st.success(f"Successfully cleaned and organized {len(final_urls)} links!")
        st.download_button(label="📥 Download Structured .docx", data=docx_data, file_name=f"Sitemap_{domain_name}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.warning("No URLs found. If using a large site, please use Tab 2 (Paste Raw XML).")
